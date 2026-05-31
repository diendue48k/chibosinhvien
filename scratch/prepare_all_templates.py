import docx
import os

public_dir = r"D:\CBSV\public"

def replace_in_runs(p, old_text, new_text):
    for r in p.runs:
        if old_text in r.text:
            r.text = r.text.replace(old_text, new_text)

def replace_in_para(p, old_text, new_text):
    if old_text in p.text:
        # If the old text is in a single run, replace it there to keep formatting
        replaced_in_run = False
        for r in p.runs:
            if old_text in r.text:
                r.text = r.text.replace(old_text, new_text)
                replaced_in_run = True
        # If not, replace in paragraph directly (might lose some run-level styling, but works)
        if not replaced_in_run:
            p.text = p.text.replace(old_text, new_text)

def replace_document_text(doc, old_text, new_text):
    # Process paragraphs
    for p in doc.paragraphs:
        replace_in_para(p, old_text, new_text)
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_para(p, old_text, new_text)

def clean_dotted_paras(doc):
    # If a paragraph consists only of dots or is completely empty after stripping, clear its text
    for p in doc.paragraphs:
        text_stripped = p.text.strip()
        if text_stripped and all(c in ' \t\n\r.….…' for c in text_stripped):
            p.text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text_stripped = p.text.strip()
                    if text_stripped and all(c in ' \t\n\r.….…' for c in text_stripped):
                        p.text = ""

# ------------------------------------------------------------
# 1. Mau 10
# ------------------------------------------------------------
print("Processing Mau 10...")
doc1 = docx.Document(os.path.join(public_dir, "1. Mau 10_KND_Ban tu kiem diem_DHKT_2025.docx"))
# Clear dotted lines first
clean_dotted_paras(doc1)
# Replacements
replacements1 = [
    ("Kính gửi: Chi uỷ ………………………………………", "Kính gửi: Chi uỷ {{chi_bo_sinh_hoat}}"),
    ("               Đảng uỷ …………………………………….", "               Đảng uỷ {{dang_uy_truong}}"),
    ("Tôi là: …………………………………….. Sinh ngày ……tháng ……năm...………", "Tôi là: {{ho_ten}} Sinh ngày {{ngay_sinh_d}} tháng {{ngay_sinh_m}} năm {{ngay_sinh_y}}"),
    ("Quê quán: \t", "Quê quán: {{que_quan}}"),
    ("Quê quán:", "Quê quán: {{que_quan}}"),
    ("+ Nơi thường trú: \t\t\t", "+ Nơi thường trú: {{dia_chi_thuong_tru}}"),
    ("+ Nơi thường trú: \t\t", "+ Nơi thường trú: {{dia_chi_thuong_tru}}"),
    ("+ Nơi thường trú: \t", "+ Nơi thường trú: {{dia_chi_thuong_tru}}"),
    ("+ Nơi thường trú:", "+ Nơi thường trú: {{dia_chi_thuong_tru}}"),
    ("+ Nơi tạm trú: \t\t\t", "+ Nơi tạm trú: {{dia_chi_tam_tru}}"),
    ("+ Nơi tạm trú: \t\t", "+ Nơi tạm trú: {{dia_chi_tam_tru}}"),
    ("+ Nơi tạm trú: \t", "+ Nơi tạm trú: {{dia_chi_tam_tru}}"),
    ("+ Nơi tạm trú:", "+ Nơi tạm trú: {{dia_chi_tam_tru}}"),
    ("Được kết nạp (hoặc kết nạp lại) vào Đảng ngày …. tháng … năm ……, tại Chi bộ …………..………………. ……………….…", "Được kết nạp (hoặc kết nạp lại) vào Đảng ngày {{ngay_vao_dang_d}} tháng {{ngay_vao_dang_m}} năm {{ngay_vao_dang_y}}, tại Chi bộ {{chi_bo_ket_nap}}"),
    ("Cơ quan, đơn vị đang công tác: .....................................................................................", "Cơ quan, đơn vị đang công tác: {{co_quan_cong_tac}}"),
    ("Đang sinh hoạt tại Chi bộ:   ...........................................................................................", "Đang sinh hoạt tại Chi bộ:   {{chi_bo_sinh_hoat}}"),
    ("Ưu điểm:………………………………………………………………………………", "Ưu điểm: {{uu_diem}}"),
    ("Khuyết điểm: ………………………………………………………………………..", "Khuyết điểm: {{khuyet_diem}}"),
    ("Biện pháp khắc phục khuyết điểm: …………………………………………………............................................................", "Biện pháp khắc phục khuyết điểm: {{bien_phap_khac_phuc}}"),
    ("………….., ngày       tháng        năm", "{{tinh_tp}}, ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}")
]
for old, new in replacements1:
    replace_document_text(doc1, old, new)
# Add student name at the end
p_last = doc1.add_paragraph()
p_last.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
p_last.add_run("\n\n\n\n\n{{ho_ten}}").bold = True
doc1.save(os.path.join(public_dir, "1. Mau 10_KND_Ban tu kiem diem_DHKT_2025.docx"))

# ------------------------------------------------------------
# 2. Nghi Quyet LCD
# ------------------------------------------------------------
print("Processing Nghi Quyet LCD...")
doc2 = docx.Document(os.path.join(public_dir, "1. Nghị quyết LCĐ.docx"))
replacements2 = [
    ("Nguyễn Hữu Ái Quốc", "{{ho_ten}}"),
    ("51K25.2", "{{lop}}"),
    ("Quản trị kinh doanh", "{{khoa}}"),
    ("QUẢN TRỊ KINH DOANH", "{{khoa_caps}}"),
    ("ngày   tháng    năm 2026", "ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}"),
    ("30 tháng 5 năm 2026", "{{ngay_vao_dang_formatted}}"),
    ("Trần Thị Lan Trinh", "{{bi_thu_lcd}}"),
    ("sự tán thành của 11 đồng chí", "sự tán thành của {{tan_thanh_lcd}} đồng chí"),
    ("đạt 100 %", "đạt {{ti_le_lcd}}%"),
    ("số không tán thành 0 đồng chí", "số không tán thành {{khong_tan_thanh_lcd}} đồng chí"),
    ("ngày    tháng     năm 2026 xét thấy:", "ngày {{ngay_hop_lcd_d}} tháng {{ngay_hop_lcd_m}} năm {{ngay_hop_lcd_y}} xét thấy:"),
    ("ngày    tháng     năm 2026", "ngày {{ngay_hop_lcd_d}} tháng {{ngay_hop_lcd_m}} năm {{ngay_hop_lcd_y}}")
]
for old, new in replacements2:
    replace_document_text(doc2, old, new)
doc2.save(os.path.join(public_dir, "1. Nghị quyết LCĐ.docx"))

# ------------------------------------------------------------
# 3. Nghi Quyet Doan Truong
# ------------------------------------------------------------
print("Processing Nghi Quyet Doan Truong...")
doc3 = docx.Document(os.path.join(public_dir, "1. Nghị quyết Đoàn Trường (02 bản).docx"))
replacements3 = [
    ("Nguyễn Hữu Ái Quốc", "{{ho_ten}}"),
    ("51K25.2", "{{lop}}"),
    ("Quản trị Kinh doanh", "{{khoa}}"),
    ("Quản trị kinh doanh", "{{khoa}}"),
    ("30 tháng 5 năm 2026", "{{ngay_vao_dang_formatted}}"),
    ("Trần Thị Lan Trinh", "{{bi_thu_doan_truong}}"),
    ("Số:       -NQ/ĐHKT-ĐTN", "Số: {{so_nq_doan_truong}}"),
    ("tán thành của 28 đồng chí", "tán thành của {{tan_thanh_doan_truong}} đồng chí"),
    ("đạt 100 %", "đạt {{ti_le_doan_truong}}%"),
    ("số không tán thành 0 đồng chí", "số không tán thành {{khong_tan_thanh_doan_truong}} đồng chí"),
    ("ngày      tháng     năm 2026", "ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}"),
    ("ngày    tháng    năm 2026 xét thấy:", "ngày {{ngay_hop_doan_truong_d}} tháng {{ngay_hop_doan_truong_m}} năm {{ngay_hop_doan_truong_y}} xét thấy:"),
    ("ngày    tháng    năm 2026", "ngày {{ngay_hop_doan_truong_d}} tháng {{ngay_hop_doan_truong_m}} năm {{ngay_hop_doan_truong_y}}")
]
for old, new in replacements3:
    replace_document_text(doc3, old, new)
doc3.save(os.path.join(public_dir, "1. Nghị quyết Đoàn Trường (02 bản).docx"))

# ------------------------------------------------------------
# 4. Bien Ban LCD
# ------------------------------------------------------------
print("Processing Bien Ban LCD...")
doc4 = docx.Document(os.path.join(public_dir, "3. Biên bản họp Liên chi Đoàn.docx"))
replacements4 = [
    ("Nguyễn Hữu Ái Quốc", "{{ho_ten}}"),
    ("51K25.2", "{{lop}}"),
    ("Quản trị kinh doanh", "{{khoa}}"),
    ("QUẢN TRỊ KINH DOANH", "{{khoa_caps}}"),
    ("Trần Thị Lan Trinh", "{{chu_tri_lcd}}"),
    ("Nguyễn Thị Xuân Hòa", "{{thu_ky_lcd}}"),
    ("11 đồng chí", "{{tong_so_uy_vien_lcd}} đồng chí"),
    ("Tham gia: 11 đồng chí", "Tham gia: {{tham_gia_lcd}} đồng chí"),
    ("Vắng: 0", "Vắng: {{vang_lcd}}"),
    ("đạt 100%", "đạt {{ti_le_lcd}}%"),
    ("số không tán thành 0 đồng chí", "số không tán thành {{khong_tan_thanh_lcd}} đồng chí"),
    ("ngày     tháng      năm 2026", "ngày {{ngay_hop_lcd_d}} tháng {{ngay_hop_lcd_m}} năm {{ngay_hop_lcd_y}}"),
    ("ngày    tháng     năm 2026", "ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}")
]
for old, new in replacements4:
    replace_document_text(doc4, old, new)
doc4.save(os.path.join(public_dir, "3. Biên bản họp Liên chi Đoàn.docx"))

# ------------------------------------------------------------
# 5. Nghi Quyet Chi Doan
# ------------------------------------------------------------
print("Processing Nghi Quyet Chi Doan...")
doc5 = docx.Document(os.path.join(public_dir, "4. Nghị quyết Chi Đoàn.docx"))
replacements5 = [
    ("Nguyễn Hữu Ái Quốc", "{{ho_ten}}"),
    ("51K25.2", "{{lop}}"),
    ("Quản trị kinh doanh", "{{khoa}}"),
    ("QUẢN TRỊ KINH DOANH", "{{khoa_caps}}"),
    ("30 tháng 5 năm 2025", "{{ngay_vao_dang_formatted}}"),
    ("ngày     tháng     năm 2026 xét thấy:", "ngày {{ngay_hop_chi_doan_d}} tháng {{ngay_hop_chi_doan_m}} năm {{ngay_hop_chi_doan_y}} xét thấy:"),
    ("ngày     tháng     năm 2026", "ngày {{ngay_hop_chi_doan_d}} tháng {{ngay_hop_chi_doan_m}} năm {{ngay_hop_chi_doan_y}}"),
    ("ngày     tháng     năm 2026", "ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}"),
    ("03 đồng chí", "{{tan_thanh_chi_doan}} đồng chí"),
    ("đạt 100 %", "đạt {{ti_le_chi_doan}}%"),
    ("số không tán thành 0 đồng chí", "số không tán thành {{khong_tan_thanh_chi_doan}} đồng chí"),
    ("Ban Chấp hành Chi Đoàn {{lop}}", "Ban Chấp hành Chi Đoàn {{lop}}"),
    ("BCH CHI ĐOÀN {{lop}}", "BCH CHI ĐOÀN {{lop}}"),
    ("TM. BAN CHẤP CHI ĐOÀN BÍ THƯ", "TM. BAN CHẤP CHI ĐOÀN\nBÍ THƯ\n\n\n\n\n{{bi_thu_chi_doan}}")
]
for old, new in replacements5:
    replace_document_text(doc5, old, new)
doc5.save(os.path.join(public_dir, "4. Nghị quyết Chi Đoàn.docx"))

# ------------------------------------------------------------
# 6. Bien Ban Chi Doan
# ------------------------------------------------------------
print("Processing Bien Ban Chi Doan...")
doc6 = docx.Document(os.path.join(public_dir, "5. Biên bản họp Chi Đoàn.docx"))
clean_dotted_paras(doc6)
replacements6 = [
    ("Chi Đoàn ………… thuộc Liên chi Đoàn khoa…..………………………...", "Chi Đoàn {{lop}} thuộc Liên chi Đoàn khoa {{khoa}}"),
    ("đồng chí …………………………….............. được kết nạp vào Đảng Cộng sản Việt Nam.", "đồng chí {{ho_ten}} được kết nạp vào Đảng Cộng sản Việt Nam."),
    ("Chi Đoàn ………… tổ chức Hội nghị Chi Đoàn để xét đề nghị công nhận đảng viên chính thức cho đoàn viên ………………………………………………", "Chi Đoàn {{lop}} tổ chức Hội nghị Chi Đoàn để xét đề nghị công nhận đảng viên chính thức cho đoàn viên {{ho_ten}}"),
    ("1. Chủ trì cuộc họp:…………………………………………………………", "1. Chủ trì cuộc họp: {{chu_tri_chi_doan}}"),
    ("2. Thư ký cuộc họp:…………………………………………………………", "2. Thư ký cuộc họp: {{thu_ky_chi_doan}}"),
    ("3. Tổng số đoàn viên: …………… Tham gia: ………… Vắng: …………..", "3. Tổng số đoàn viên: {{tong_so_dv_chi_doan}}   Tham gia: {{tham_gia_chi_doan}}   Vắng: {{vang_chi_doan}}"),
    ("Lý do:………………………………………………………………………...", "Lý do vắng: {{ly_do_vang_chi_doan}}"),
    ("Chi Đoàn …………… đã tiến hành họp nhận xét thông qua bản tự kiểm điểm của đồng chí ………………………………… trong suốt thời gian dự bị.", "Chi Đoàn {{lop}} đã tiến hành họp nhận xét thông qua bản tự kiểm điểm của đồng chí {{ho_ten}} trong suốt thời gian dự bị."),
    ("Hội nghị nhất trí nhận xét về những ưu điểm, khuyết điểm chính của đồng chí ………………………………………. trong suốt quá trình phấn đấu học tập, công tác, rèn luyện tại Chi Đoàn, cụ thể như sau:", "Hội nghị nhất trí nhận xét về những ưu điểm, khuyết điểm chính của đồng chí {{ho_ten}} trong suốt quá trình phấn đấu học tập, công tác, rèn luyện tại Chi Đoàn, cụ thể như sau:"),
    ("1. Ưu điểm:", "1. Ưu điểm: {{uu_diem}}"),
    ("2. Khuyết điểm:", "2. Khuyết điểm: {{khuyet_diem}}"),
    ("Chi Đoàn…………….. thuộc Liên chi Đoàn khoa …………………………………… giới thiệu", "Chi Đoàn {{lop}} thuộc Liên chi Đoàn khoa {{khoa}} giới thiệu"),
    ("đồng chí ……………………………………………", "đồng chí {{ho_ten}}"),
    ("tán thành của …… đoàn viên (đạt ……%); số không tán thành …… đồng chí, với lý do………………………………………….", "tán thành của {{tham_gia_chi_doan}} đoàn viên (đạt {{ti_le_bb_chi_doan}}%); số không tán thành {{khong_tan_thanh_bb_chi_doan}} đồng chí, với lý do: {{ly_do_khong_tan_thanh_bb_chi_doan}}"),
    ("Tập thể Chi Đoàn ……… thuộc Liên chi Đoàn khoa………………………...", "Tập thể Chi Đoàn {{lop}} thuộc Liên chi Đoàn khoa {{khoa}}"),
    ("ngày …… tháng …… năm …… tại Trường", "ngày {{ngay_hop_chi_doan_d}} tháng {{ngay_hop_chi_doan_m}} năm {{ngay_hop_chi_doan_y}}"),
    ("ngày……tháng…… năm 20…", "ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}"),
    ("LIÊN CHI ĐOÀN KHOA …………………………………….. BCH CHI ĐOÀN…………………", "LIÊN CHI ĐOÀN KHOA {{khoa_caps}} BCH CHI ĐOÀN {{lop}}"),
    ("Thư ký (Ký, ghi họ tên)", "Thư ký\n\n\n\n\n{{thu_ky_chi_doan}}"),
    ("Chủ tọa (Ký, ghi họ tên)", "Chủ tọa\n\n\n\n\n{{chu_tri_chi_doan}}")
]
for old, new in replacements6:
    replace_document_text(doc6, old, new)
doc6.save(os.path.join(public_dir, "5. Biên bản họp Chi Đoàn.docx"))

# ------------------------------------------------------------
# 7. Bien Ban Hop Lop
# ------------------------------------------------------------
print("Processing Bien Ban Hop Lop...")
doc7 = docx.Document(os.path.join(public_dir, "6. Biên bản họp lớp.docx"))
clean_dotted_paras(doc7)
replacements7 = [
    ("ngày …… tháng …… năm …… tại Trường", "ngày {{ngay_hop_lop_d}} tháng {{ngay_hop_lop_m}} năm {{ngay_hop_lop_y}}"),
    ("Lớp ………… thuộc khoa ………………………...", "Lớp {{lop}} thuộc khoa {{khoa}}"),
    ("cho sinh viên ……………………………………………………………………………….", "cho sinh viên {{ho_ten}}"),
    ("1. Giảng viên chủ nhiệm:…………………………………………………...", "1. Giảng viên chủ nhiệm: {{gvcn}}"),
    ("2. Chủ trì cuộc họp:…………………………………………………………", "2. Chủ trì cuộc họp: {{chu_tri_lop}}"),
    ("3. Thư ký cuộc họp:…………………………………………………………", "3. Thư ký cuộc họp: {{thu_ky_lop}}"),
    ("4. Tổng số sinh viên lớp: ………… Tham gia: ………… Vắng: …………", "4. Tổng số sinh viên lớp: {{tong_so_sv_lop}}   Tham gia: {{tham_gia_lop}}   Vắng: {{vang_lop}}"),
    ("đảng viên dự bị ………………………………………. trong suốt quá trình phấn đấu học tập, công tác, rèn luyện tại lớp, cụ thể như sau:", "đảng viên dự bị {{ho_ten}} trong suốt quá trình phấn đấu học tập, công tác, rèn luyện tại lớp, cụ thể như sau:"),
    ("1. Ưu điểm:", "1. Ưu điểm: {{uu_diem}}"),
    ("2. Khuyết điểm:", "2. Khuyết điểm: {{khuyet_diem}}"),
    ("tập thể lớp………… thuộc khoa …………………………… giới thiệu", "tập thể lớp {{lop}} thuộc khoa {{khoa}} giới thiệu"),
    ("cho đảng viên dự bị………………………………………………………………..", "cho đảng viên dự bị {{ho_ten}}"),
    ("tán thành của …… sinh viên (đạt ……%); số không tán thành …… sinh viên, với lý do…………………………………………", "tán thành của {{tham_gia_lop}} sinh viên (đạt {{ti_le_bb_lop}}%); số không tán thành {{khong_tan_thanh_bb_lop}} sinh viên, với lý do: {{ly_do_khong_tan_thanh_bb_lop}}"),
    ("Tập thể lớp………… thuộc khoa………………………………. xin chịu", "Tập thể lớp {{lop}} thuộc khoa {{khoa}} xin chịu"),
    ("Giảng viên chủ nhiệm lớp (Ký, ghi họ tên)", "Giảng viên chủ nhiệm lớp\n\n\n\n\n{{gvcn}}"),
    ("Thư ký (Ký, ghi họ tên)", "Thư ký\n\n\n\n\n{{thu_ky_lop}}"),
    ("Chủ tọa (Ký, ghi họ tên)", "Chủ tọa\n\n\n\n\n{{chu_tri_lop}}"),
    ("ngày……tháng… năm 20…", "ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}")
]
for old, new in replacements7:
    replace_document_text(doc7, old, new)
doc7.save(os.path.join(public_dir, "6. Biên bản họp lớp.docx"))

print("All templates successfully prepared with placeholders!")
