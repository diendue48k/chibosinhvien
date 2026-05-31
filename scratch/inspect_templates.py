import docx
import os
import sys

# Force stdout to be utf-8
sys.stdout.reconfigure(encoding='utf-8')

templates = [
    "1. Mau 10_KND_Ban tu kiem diem_DHKT_2025.doc",
    "1. Nghị quyết LCĐ.docx",
    "1. Nghị quyết Đoàn Trường (02 bản).docx",
    "3. Biên bản họp Liên chi Đoàn.docx",
    "4. Nghị quyết Chi Đoàn.docx",
    "5. Biên bản họp Chi Đoàn.docx",
    "6. Biên bản họp lớp.docx",
    "ĐHKT_QĐ 213- 24052026.Sinh vien.docx"
]

public_dir = r"D:\CBSV\public"

print("--- INSPECTING TEMPLATES ---")
for t in templates:
    path = os.path.join(public_dir, t)
    if not os.path.exists(path):
        print(f"File not found: {t}")
        continue
    
    print(f"\n==================== {t} ====================")
    if t.endswith('.doc'):
        print(f"Skipping binary doc file (requires special parser), size: {os.path.getsize(path)} bytes")
        continue
        
    try:
        doc = docx.Document(path)
        print("Paragraphs:")
        p_count = 0
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                print(f"  P: {text}")
                p_count += 1
                if p_count >= 10:
                    break
        
        print("Tables:")
        for tbl_idx, table in enumerate(doc.tables):
            print(f"  Table {tbl_idx}:")
            for row_idx, row in enumerate(table.rows[:4]):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.replace("\n", " ").strip()
                    row_text.append(cell_text[:50])
                print(f"    Row {row_idx}: {row_text}")
    except Exception as e:
        print(f"Error reading {t}: {e}")
