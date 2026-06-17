import docx
import os

public_dir = r"D:\CBSV\public"

def replace_in_para(p, old_text, new_text):
    if old_text in p.text:
        replaced_in_run = False
        for r in p.runs:
            if old_text in r.text:
                r.text = r.text.replace(old_text, new_text)
                replaced_in_run = True
        if not replaced_in_run:
            p.text = p.text.replace(old_text, new_text)

def replace_document_text(doc, old_text, new_text):
    # Process paragraphs
    for p in doc.paragraphs:
        replace_in_para(p, old_text, new_text)
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_para(p, old_text, new_text)

# ------------------------------------------------------------
# 1. Mau 11
# ------------------------------------------------------------
print("Processing Mau 11...")
doc11 = docx.Document(os.path.join(public_dir, "2. Mau 11_KND_nhan xet dang vien giup do_DHKT_2025.docx"))

replacements11 = [
    ("Nguyễn Hữu Ái Quốc", "{{ho_ten}}"),
    ("Lê Vĩnh Diện", "{{dvhd}}"),
    ("Đà Nẵng, ngày        tháng       năm 2026", "{{tinh_tp}}, ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}"),
    ("- Chi ủy Chi bộ Sinh viên", "- Chi ủy Chi bộ {{chi_bo_sinh_hoat}}"),
    ("- Đảng ủy Trường Đại học Kinh tế", "- Đảng ủy {{dang_uy_truong}}"),
    ("01/04/2004", "{{dvhd_ngay_sinh_formatted}}"),
    ("Ngày vào Đảng: 26/07/2022", "Ngày vào Đảng: {{dvhd_ngay_vao_dang}}"),
    ("Ngày chính thức: 26/07/2023", "Ngày chính thức: {{dvhd_ngay_chinh_thuc}}"),
    ("Hiện đang sinh hoạt tại Chi bộ Sinh viên", "Hiện đang sinh hoạt tại Chi bộ {{chi_bo_sinh_hoat}}"),
    ("Ngày      tháng       năm 2025", "Ngày {{ngay_phan_cong_d}} tháng {{ngay_phan_cong_m}} năm {{ngay_phan_cong_y}}"),
    ("ngày 30 tháng 5 năm 2025", "ngày {{ngay_vao_dang_d}} tháng {{ngay_vao_dang_m}} năm {{ngay_vao_dang_y}}")
]

for old, new in replacements11:
    replace_document_text(doc11, old, new)
doc11.save(os.path.join(public_dir, "2. Mau 11_KND_nhan xet dang vien giup do_DHKT_2025.docx"))

# ------------------------------------------------------------
# 2. Mau 12
# ------------------------------------------------------------
print("Processing Mau 12...")
doc12 = docx.Document(os.path.join(public_dir, "7. Mau 12_KND_Tong hop nhan xet cuadoan the dv dang vien du bi.docx"))

replacements12 = [
    ("Nguyễn Hữu Ái Quốc", "{{ho_ten}}"),
    ("30 tháng 5 năm 2025", "{{ngay_vao_dang_formatted_vietnamese}}"),
    ("Khoa Quản trị kinh doanh, chi Đoàn 51K25.2", "Khoa {{khoa}}, chi Đoàn {{lop}}"),
    ("tổng số có 100 đồng chí", "tổng số có {{tong_so_to_chuc_ctxh}} đồng chí"),
    ("43-44-45 An Thượng, có 3 đồng chí", "{{chi_uy_noi_cu_tru}}, có {{tong_so_chi_uy_noi_cu_tru}} đồng chí"),
    ("là 100 đồng chí, trong tổng số 100 đồng chí được hỏi ý kiến (đạt 100%)", "là {{tong_so_to_chuc_ctxh}} đồng chí, trong tổng số {{tong_so_to_chuc_ctxh}} đồng chí được hỏi ý kiến (đạt 100%)"),
    ("Số không tán thành 0 đồng chí", "Số không tán thành {{khong_tan_thanh_ctxh}} đồng chí"),
    ("là 3 đồng chí, trong tổng số 3 đồng chí được hỏi ý kiến (đạt 100%)", "là {{tong_so_chi_uy_noi_cu_tru}} đồng chí, trong tổng số {{tong_so_chi_uy_noi_cu_tru}} đồng chí được hỏi ý kiến (đạt 100%)"),
    ("TRƯỜNG ĐẠI HỌC KINH TẾ", "{{dang_uy_truong_caps}}"),
    ("CHI BỘ SINH VIÊN", "CHI BỘ {{chi_bo_sinh_hoat_caps}}"),
    ("Đà Nẵng, ngày      tháng     năm 2026", "{{tinh_tp}}, ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}"),
    ("Bùi Trung Hiệp", "{{chu_tri_chi_bo}}")
]

for old, new in replacements12:
    replace_document_text(doc12, old, new)
doc12.save(os.path.join(public_dir, "7. Mau 12_KND_Tong hop nhan xet cuadoan the dv dang vien du bi.docx"))

# ------------------------------------------------------------
# 3. Mau 13
# ------------------------------------------------------------
print("Processing Mau 13...")
doc13 = docx.Document(os.path.join(public_dir, "8. Mau 13_KND_Nghi quyet de nghi chinh thuc chi bo.docx"))

replacements13 = [
    ("TRƯỜNG ĐẠI HỌC KINH TẾ", "{{dang_uy_truong_caps}}"),
    ("CHI BỘ SINH VIÊN", "CHI BỘ {{chi_bo_sinh_hoat_caps}}"),
    ("Đà Nẵng, ngày    tháng    năm 2026", "{{tinh_tp}}, ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}"),
    ("Ngày      tháng      năm 2026", "Ngày {{ngay_hop_chi_bo_d}} tháng {{ngay_hop_chi_bo_m}} năm {{ngay_hop_chi_bo_y}}"),
    ("Chi bộ Sinh viên đã họp", "Chi bộ {{chi_bo_sinh_hoat}} đã họp"),
    ("Nguyễn Hữu Ái Quốc", "{{ho_ten}}"),
    ("ngày 30 tháng 5 năm 2026", "ngày {{ngay_vao_dang_d}} tháng {{ngay_vao_dang_m}} năm {{ngay_vao_dang_y}}"),
    ("ngày 30 tháng 5 năm 2025", "ngày {{ngay_vao_dang_d}} tháng {{ngay_vao_dang_m}} năm {{ngay_vao_dang_y}}"),
    ("378 đảng viên, trong đó chính thức 234 đồng chí, dự bị 144 đồng chí", "{{tong_so_dv}} đảng viên, trong đó chính thức {{tong_so_dv_chinh_thuc}} đồng chí, dự bị {{tong_so_dv_du_bi}} đồng chí"),
    ("Vắng mặt: 0 đảng viên, trong đó chính thức 0 đồng chí, dự bị 0 đồng chí", "Vắng mặt: {{vang_chi_bo}} đảng viên, trong đó chính thức {{vang_chinh_thuc_chi_bo}} đồng chí, dự bị {{vang_du_bi_chi_bo}} đồng chí"),
    ("Bùi Trung Hiệp", "{{chu_tri_chi_bo}}"),
    ("Bí thư Chi bộ", "{{chuc_vu_chu_tri_chi_bo}}"),
    ("Lê Vĩnh Diện", "{{thu_ky_chi_bo}}"),
    ("là 234 đồng chí (đạt 100%)", "là {{tan_thanh_chi_bo}} đồng chí (đạt {{ti_le_chi_bo}}%)"),
    ("không tán thành 0 đồng chí (chiếm 0%)", "không tán thành {{khong_tan_thanh_chi_bo}} đồng chí (chiếm {{ti_le_khong_tan_thanh_chi_bo}}%)"),
    ("Đảng uỷ Trường ĐHKT(để b/c)", "Đảng uỷ {{dang_uy_truong}}(để b/c)")
]

for old, new in replacements13:
    replace_document_text(doc13, old, new)
doc13.save(os.path.join(public_dir, "8. Mau 13_KND_Nghi quyet de nghi chinh thuc chi bo.docx"))

print("All Chi bo templates successfully processed!")
