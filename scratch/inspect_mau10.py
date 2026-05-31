import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

path = r"D:\CBSV\public\1. Mau 10_KND_Ban tu kiem diem_DHKT_2025.docx"
try:
    doc = docx.Document(path)
    print("--- PARAGRAPHS FOR MAU 10 ---")
    for idx, p in enumerate(doc.paragraphs):
        print(f"P{idx}: {p.text}")
    print("--- TABLES FOR MAU 10 ---")
    for t_idx, table in enumerate(doc.tables):
        print(f"Table {t_idx}:")
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            print(f"  Row {r_idx}: {cells}")
except Exception as e:
    print("Error:", e)
