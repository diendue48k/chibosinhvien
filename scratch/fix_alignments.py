import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

public_dir = r"D:\CBSV\public"

files_to_fix = [
    "1. Mau 10_KND_Ban tu kiem diem_DHKT_2025.docx",
    "5. Biên bản họp Chi Đoàn.docx",
    "6. Biên bản họp lớp.docx"
]

print("--- FIXING PARAGRAPH ALIGNMENTS ---")
for f_name in files_to_fix:
    path = os.path.join(public_dir, f_name)
    if not os.path.exists(path):
        print(f"File not found: {f_name}")
        continue
    
    doc = docx.Document(path)
    fixed_count = 0
    
    # Process paragraphs
    for p in doc.paragraphs:
        if "{{uu_diem}}" in p.text or "{{khuyet_diem}}" in p.text or "{{bien_phap_khac_phuc}}" in p.text:
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            fixed_count += 1
            print(f"  Fixed alignment for paragraph: '{p.text[:40]}...' in {f_name}")
            
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "{{uu_diem}}" in p.text or "{{khuyet_diem}}" in p.text or "{{bien_phap_khac_phuc}}" in p.text:
                        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                        fixed_count += 1
                        print(f"  Fixed alignment for table cell: '{p.text[:40]}...' in {f_name}")
                        
    doc.save(path)
    print(f"Completed {f_name}, fixed {fixed_count} paragraphs.")
