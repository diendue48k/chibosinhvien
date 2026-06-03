import docx
import os
import sys

# Ensure UTF-8 stdout to handle Vietnamese accents in console print
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

public_dir = r"D:\CBSV\public"
files_to_fix = [
    "3. Nghị quyết LCĐ.docx",
    "2. Nghị quyết Đoàn Trường (02 bản).docx",
    "3. Biên bản họp Liên chi Đoàn.docx",
    "4. Nghị quyết Chi Đoàn.docx"
]

def fix_template(filename):
    filepath = os.path.join(public_dir, filename)
    print(f"Fixing: {filename}")
    doc = docx.Document(filepath)
    
    paragraphs_to_remove = []
    in_uu_diem = False
    in_khuyet_diem = False
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("1. Ưu điểm:"):
            p.text = "1. Ưu điểm: {{uu_diem}}"
            in_uu_diem = True
            in_khuyet_diem = False
            continue
        elif text.startswith("2. Khuyết điểm:"):
            p.text = "2. Khuyết điểm: {{khuyet_diem}}"
            in_uu_diem = False
            in_khuyet_diem = True
            continue
        elif in_uu_diem:
            if text.startswith("-"):
                paragraphs_to_remove.append(p)
            else:
                in_uu_diem = False
        elif in_khuyet_diem:
            if text.startswith("-"):
                paragraphs_to_remove.append(p)
            else:
                in_khuyet_diem = False
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                in_uu_diem = False
                in_khuyet_diem = False
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text.startswith("1. Ưu điểm:"):
                        p.text = "1. Ưu điểm: {{uu_diem}}"
                        in_uu_diem = True
                        in_khuyet_diem = False
                        continue
                    elif text.startswith("2. Khuyết điểm:"):
                        p.text = "2. Khuyết điểm: {{khuyet_diem}}"
                        in_uu_diem = False
                        in_khuyet_diem = True
                        continue
                    elif in_uu_diem:
                        if text.startswith("-"):
                            paragraphs_to_remove.append(p)
                        else:
                            in_uu_diem = False
                    elif in_khuyet_diem:
                        if text.startswith("-"):
                            paragraphs_to_remove.append(p)
                        else:
                            in_khuyet_diem = False

    # Remove marked paragraphs from document
    for p in paragraphs_to_remove:
        pElement = p._element
        pElement.getparent().remove(pElement)
        
    doc.save(filepath)
    print(f"Successfully fixed and saved {filename}")

for filename in files_to_fix:
    fix_template(filename)
