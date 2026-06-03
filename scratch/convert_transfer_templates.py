import docx
import os
import sys

# Ensure UTF-8 stdout
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

public_dir = r"D:\CBSV\public\CHUYỂN SINH HOẠT"

# Helper for paragraph replacement
def replace_para_if_contains(p, search_str, new_text):
    if search_str in p.text:
        p.text = new_text
        return True
    return False

# ------------------------------------------------------------
# 1. Don xin chuyen dang
# ------------------------------------------------------------
print("Processing Don xin chuyen dang...")
doc1 = docx.Document(os.path.join(public_dir, "1. Mẫu 1. Don xin chuyen dang.docx"))
for p in doc1.paragraphs:
    replace_para_if_contains(p, "Đà Nẵng, ngày    tháng     năm 2026", "{{tinh_tp}}, ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}")
    replace_para_if_contains(p, "Họ và tên : Lê Vĩnh Diện Giới tính: Nam", "Họ và tên : {{ho_ten}} Giới tính: {{gioi_tinh}}")
    replace_para_if_contains(p, "Ngày tháng năm sinh: 01/04/2004", "Ngày tháng năm sinh: {{ngay_sinh_formatted}}")
    replace_para_if_contains(p, "Quê quán:  Xã Vĩnh Thủy, tỉnh Quảng Trị", "Quê quán: {{que_quan}}")
    replace_para_if_contains(p, "Nơi ở hiện nay: Thôn Đức Xá, Xã Vĩnh Thủy, tỉnh Quảng Trị", "Nơi ở hiện nay: {{dia_chi}}")
    replace_para_if_contains(p, "Ngày vào Đảng: 26/07/2022", "Ngày vào Đảng: {{ngay_vao_dang_formatted}}  Ngày chính thức: {{ngay_chinh_thuc_formatted}}")
    replace_para_if_contains(p, "Số thẻ Đảng: 045204008389", "Số thẻ Đảng: {{so_the_dang}}  Số điện thoại liên hệ: {{so_dien_thoai}}")
    replace_para_if_contains(p, "Nhiệm vụ được giao:  Đảng viên, Phó bí thư Chi bộ Sinh viên", "Nhiệm vụ được giao: {{nhiem_vu_dang}}")
    replace_para_if_contains(p, "Tôi đã hoàn thành chương trình học và đã tốt nghiệp ra trường", "{{ly_do_chuyen}}")
    replace_para_if_contains(p, "Để thực hiện theo đúng Điều lệ Đảng, kính đề nghị các cấp ủy Đảng cho tôi được chuyển sinh hoạt Đảng về Chi bộ trực thuộc thôn Đức Xá thuộc Đảng bộ cơ sở xã Vĩnh Thủy, Đảng bộ cấp trên cơ sở tỉnh Quảng Trị", "Để thực hiện theo đúng Điều lệ Đảng, kính đề nghị các cấp ủy Đảng cho tôi được chuyển sinh hoạt Đảng về {{noi_chuyen_den}}")

# Fix signature at the end
for p in reversed(doc1.paragraphs):
    if p.text.strip() == "Lê Vĩnh Diện":
        p.text = "{{ho_ten}}"
        break

doc1.save(os.path.join(public_dir, "1. Mẫu 1. Don xin chuyen dang.docx"))

# ------------------------------------------------------------
# 2. Don xin chuyen dang tam thoi
# ------------------------------------------------------------
print("Processing Don xin chuyen dang tam thoi...")
doc2 = docx.Document(os.path.join(public_dir, "2. Mau 2. Don xin chuyen dang tam thoi.docx"))
for p in doc2.paragraphs:
    replace_para_if_contains(p, "Đà Nẵng, ngày    tháng     năm 2026", "{{tinh_tp}}, ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}")
    replace_para_if_contains(p, "Họ và tên : Lê Vĩnh Diện Giới tính: Nam", "Họ và tên : {{ho_ten}} Giới tính: {{gioi_tinh}}")
    replace_para_if_contains(p, "Ngày tháng năm sinh: 01/04/2004", "Ngày tháng năm sinh: {{ngay_sinh_formatted}}")
    replace_para_if_contains(p, "Quê quán:  Xã Vĩnh Thủy, tỉnh Quảng Trị", "Quê quán: {{que_quan}}")
    replace_para_if_contains(p, "Nơi ở hiện nay: Thôn Đức Xá, Xã Vĩnh Thủy, tỉnh Quảng Trị", "Nơi ở hiện nay: {{dia_chi}}")
    replace_para_if_contains(p, "Ngày vào Đảng: 26/07/2022", "Ngày vào Đảng: {{ngay_vao_dang_formatted}}  Ngày chính thức: {{ngay_chinh_thuc_formatted}}")
    replace_para_if_contains(p, "Số thẻ Đảng: 045204008389", "Số thẻ Đảng: {{so_the_dang}}  Số điện thoại liên hệ: {{so_dien_thoai}}")
    replace_para_if_contains(p, "Nhiệm vụ được giao:  Đảng viên, Phó bí thư Chi bộ Sinh viên", "Nhiệm vụ được giao: {{nhiem_vu_dang}}")
    replace_para_if_contains(p, "Tôi hiện đang đi thực tập ở quê nên cần phải chuyển sinh hoạt tạm thời", "{{ly_do_chuyen}}")
    replace_para_if_contains(p, "Để thực hiện theo đúng Điều lệ Đảng, kính đề nghị các cấp ủy Đảng cho tôi được chuyển sinh hoạt Đảng về Chi bộ trực thuộc thôn Đức Xá thuộc Đảng bộ cơ sở xã Vĩnh Thủy, Đảng bộ cấp trên cơ sở tỉnh Quảng Trị", "Để thực hiện theo đúng Điều lệ Đảng, kính đề nghị các cấp ủy Đảng cho tôi được chuyển sinh hoạt Đảng về {{noi_chuyen_den}}")

for p in reversed(doc2.paragraphs):
    if p.text.strip() == "Lê Vĩnh Diện":
        p.text = "{{ho_ten}}"
        break

doc2.save(os.path.join(public_dir, "2. Mau 2. Don xin chuyen dang tam thoi.docx"))

# ------------------------------------------------------------
# 3. Bản nhận xét Đảng viên dự bị ĐTN
# ------------------------------------------------------------
print("Processing Nhan xet ĐTN...")
doc3 = docx.Document(os.path.join(public_dir, "3. Mẫu 3. Bản nhận xét Đảng viên dự bị ĐTN.docx"))
for p in doc3.paragraphs:
    replace_para_if_contains(p, "Số         - NQ/ĐTN-ĐHKT", "Số: {{so_nq_doan_truong}}")
    replace_para_if_contains(p, "Đà Nẵng, ngày     tháng      năm 2026", "{{tinh_tp}}, ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}")
    replace_para_if_contains(p, "Đảng viên dự bị: Nguyễn Văn An", "Đảng viên dự bị: {{ho_ten}}          Sinh ngày: {{ngay_sinh_formatted}}")
    replace_para_if_contains(p, "Chi Đoàn: 48K27", "Chi Đoàn: {{lop}}          Liên chi Đoàn Khoa: {{khoa}}")
    replace_para_if_contains(p, "Được kết nạp vào Đảng ngày 26 tháng 05 năm 2026.", "Được kết nạp vào Đảng ngày {{ngay_vao_dang_formatted_vietnamese}}.")

# Replace advantages and disadvantages
paragraphs_to_remove = []
in_uu_diem = False
in_khuyet_diem = False
for p in doc3.paragraphs:
    text = p.text.strip()
    if text == "Ưu điểm:":
        p.text = "Ưu điểm:\n{{uu_diem}}"
        in_uu_diem = True
        in_khuyet_diem = False
    elif text == "Khuyết điểm và những vấn đề cần lưu ý:":
        p.text = "Khuyết điểm và những vấn đề cần lưu ý:\n{{khuyet_diem}}"
        in_uu_diem = False
        in_khuyet_diem = True
    elif in_uu_diem:
        if text.startswith("-"):
            paragraphs_to_remove.append(p)
        else:
            in_uu_diem = False
    elif in_khuyet_diem:
        if text.startswith("-"):
            paragraphs_to_remove.append(p)
        else:
            in_khuyet_diem = False

for p in paragraphs_to_remove:
    pElement = p._element
    pElement.getparent().remove(pElement)

doc3.save(os.path.join(public_dir, "3. Mẫu 3. Bản nhận xét Đảng viên dự bị ĐTN.docx"))

# ------------------------------------------------------------
# 4. Bản nhận xét Đảng viên dự bị Chuyển SHĐ ĐVHD
# ------------------------------------------------------------
print("Processing Nhan xet ĐVHD...")
doc4 = docx.Document(os.path.join(public_dir, "4. Mẫu 5. Bản nhận xét Đảng viên dự bị Chuyển SHĐ ĐVHD.docx"))
for p in doc4.paragraphs:
    replace_para_if_contains(p, "Tôi là: Lê Vĩnh Diện", "Tôi là: {{dvhd}}          Sinh ngày {{dvhd_ngay_sinh_formatted}}")
    replace_para_if_contains(p, "Ngày vào Đảng: 26/07/2022", "Ngày vào Đảng: {{dvhd_ngay_vao_dang}}          Chính thức: {{dvhd_ngay_chinh_thuc}}")
    replace_para_if_contains(p, "Ngày 26 tháng 07 năm 2026 được Chi bộ phân công giúp đỡ", "Ngày {{ngay_phan_cong}} được Chi bộ phân công giúp đỡ đảng viên dự bị: {{ho_ten}} được kết nạp vào Đảng ngày {{ngay_vao_dang_formatted_vietnamese}}, nay xin báo cáo Đảng ủy và Chi bộ những vấn đề chủ yếu của đảng viên dự bị như sau:")
    replace_para_if_contains(p, "Đà Nẵng, ngày       tháng         năm 2026", "{{tinh_tp}}, ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}")

# Replace advantages and disadvantages
paragraphs_to_remove = []
in_uu_diem = False
in_khuyet_diem = False
for p in doc4.paragraphs:
    text = p.text.strip()
    if text == "Ưu điểm:":
        p.text = "Ưu điểm:\n{{uu_diem}}"
        in_uu_diem = True
        in_khuyet_diem = False
    elif text == "Khuyết điểm và những vấn đề cần lưu ý:":
        p.text = "Khuyết điểm và những vấn đề cần lưu ý:\n{{khuyet_diem}}"
        in_uu_diem = False
        in_khuyet_diem = True
    elif in_uu_diem:
        if text.startswith("-"):
            paragraphs_to_remove.append(p)
        else:
            in_uu_diem = False
    elif in_khuyet_diem:
        if text.startswith("-"):
            paragraphs_to_remove.append(p)
        else:
            in_khuyet_diem = False

for p in paragraphs_to_remove:
    pElement = p._element
    pElement.getparent().remove(pElement)

for p in reversed(doc4.paragraphs):
    if p.text.strip() == "Lê Vĩnh Diện":
        p.text = "{{dvhd}}"
        break

doc4.save(os.path.join(public_dir, "4. Mẫu 5. Bản nhận xét Đảng viên dự bị Chuyển SHĐ ĐVHD.docx"))

# ------------------------------------------------------------
# 5. Kiem diem chuyen dang 2026
# ------------------------------------------------------------
print("Processing Kiem diem chuyen dang...")
doc5 = docx.Document(os.path.join(public_dir, "5. Mẫu 4. Kiem diem chuyen dang 2026.docx"))
for p in doc5.paragraphs:
    replace_para_if_contains(p, "Đà Nẵng, ngày    tháng     năm 2026", "{{tinh_tp}}, ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}")
    replace_para_if_contains(p, "Họ và tên: Lê Vĩnh Diện Năm sinh: 2004", "Họ và tên: {{ho_ten}}          Năm sinh: {{nam_sinh}}")
    replace_para_if_contains(p, "Ngày vào Đảng: 26/07/2022", "Ngày vào Đảng: {{ngay_vao_dang_formatted}}          Ngày chính thức: {{ngay_chinh_thuc_formatted}}     Chi bộ đang sinh hoạt: {{chi_bo_sinh_hoat}}")
    replace_para_if_contains(p, "+ Đảng: Đảng viên, Phó Bí thư Chi bộ Sinh viên", "+ Đảng: {{nhiem_vu_dang}}")
    replace_para_if_contains(p, "Đà Nẵng, ngày         tháng        năm 2026", "{{tinh_tp}}, ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}")

for p in reversed(doc5.paragraphs):
    if p.text.strip() == "Lê Vĩnh Diện":
        p.text = "{{ho_ten}}"
        break

doc5.save(os.path.join(public_dir, "5. Mẫu 4. Kiem diem chuyen dang 2026.docx"))

print("All transfer-out templates successfully placeholder-ized!")
