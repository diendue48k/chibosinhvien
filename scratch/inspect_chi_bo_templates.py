import docx
import os
import sys

# Reconfigure stdout/stderr if possible
try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

public_dir = r"D:\CBSV\public"

files = [
    "2. Mau 11_KND_nhan xet dang vien giup do_DHKT_2025.docx",
    "7. Mau 12_KND_Tong hop nhan xet cuadoan the dv dang vien du bi.docx",
    "8. Mau 13_KND_Nghi quyet de nghi chinh thuc chi bo.docx"
]

with open(r"d:\CBSV\scratch\chi_bo_templates_dump.txt", "w", encoding="utf-8") as out:
    for filename in files:
        path = os.path.join(public_dir, filename)
        out.write("=" * 60 + "\n")
        out.write(f"FILE: {filename}\n")
        out.write("=" * 60 + "\n")
        doc = docx.Document(path)
        
        # Print all paragraphs
        for idx, p in enumerate(doc.paragraphs):
            if p.text.strip():
                out.write(f"P{idx}: {p.text}\n")
                
        # Also check tables
        for t_idx, table in enumerate(doc.tables):
            out.write(f"--- Table {t_idx} ---\n")
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, p in enumerate(cell.paragraphs):
                        if p.text.strip():
                            out.write(f"  T{t_idx} R{r_idx} C{c_idx} P{p_idx}: {p.text}\n")
print("Done dumping!")
