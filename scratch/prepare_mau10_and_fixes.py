import docx
import os

public_dir = r"D:\CBSV\public"

# ------------------------------------------------------------
# 1. Process pristine Mau 10 with safe paragraph-level replacements
# ------------------------------------------------------------
print("Processing Mau 10...")
doc1 = docx.Document(os.path.join(public_dir, "1. Mau 10_KND_Ban tu kiem diem_DHKT_2025.docx"))

# Helper for paragraph replacement
def replace_para_if_starts_with(p, prefix, new_text):
    if p.text.strip().startswith(prefix):
        p.text = new_text
        return True
    return False

def replace_para_if_contains(p, search_str, new_text):
    if search_str in p.text:
        p.text = new_text
        return True
    return False

# Clear paragraphs consisting only of dots
for p in doc1.paragraphs:
    stripped = p.text.strip()
    if stripped and all(c in ' \t\n\r.….…' for c in stripped):
        p.text = ""

for p in doc1.paragraphs:
    # Safely replace full paragraphs to avoid overlapping or duplications
    if replace_para_if_starts_with(p, "Kính gửi: Chi uỷ", "Kính gửi: Chi uỷ {{chi_bo_sinh_hoat}}"):
        continue
    if replace_para_if_starts_with(p, "               Đảng uỷ", "               Đảng uỷ {{dang_uy_truong}}"):
        continue
    if replace_para_if_starts_with(p, "Tôi là: ……………………………………..", "Tôi là: {{ho_ten}}          Sinh ngày {{ngay_sinh_d}} tháng {{ngay_sinh_m}} năm {{ngay_sinh_y}}"):
        continue
    if replace_para_if_starts_with(p, "Quê quán:", "Quê quán: {{que_quan}}"):
        continue
    if replace_para_if_starts_with(p, "+ Nơi thường trú:", "+ Nơi thường trú: {{dia_chi_thuong_tru}}"):
        continue
    if replace_para_if_starts_with(p, "+ Nơi tạm trú:", "+ Nơi tạm trú: {{dia_chi_tam_tru}}"):
        continue
    if replace_para_if_starts_with(p, "Được kết nạp", "Được kết nạp (hoặc kết nạp lại) vào Đảng ngày {{ngay_vao_dang_d}} tháng {{ngay_vao_dang_m}} năm {{ngay_vao_dang_y}}, tại Chi bộ {{chi_bo_ket_nap}}"):
        continue
    if replace_para_if_starts_with(p, "Cơ quan, đơn vị đang công tác:", "Cơ quan, đơn vị đang công tác: {{co_quan_cong_tac}}"):
        continue
    if replace_para_if_starts_with(p, "Đang sinh hoạt tại Chi bộ:", "Đang sinh hoạt tại Chi bộ: {{chi_bo_sinh_hoat}}"):
        continue
    if replace_para_if_starts_with(p, "Ưu điểm:", "Ưu điểm: {{uu_diem}}"):
        continue
    if replace_para_if_starts_with(p, "Khuyết điểm:", "Khuyết điểm: {{khuyet_diem}}"):
        continue
    if replace_para_if_starts_with(p, "Biện pháp khắc phục khuyết điểm:", "Biện pháp khắc phục khuyết điểm: {{bien_phap_khac_phuc}}"):
        continue
    if replace_para_if_contains(p, "ngày       tháng        năm", "\t\t\t\t\t{{tinh_tp}}, ngày {{ngay_ky_d}} tháng {{ngay_ky_m}} năm {{ngay_ky_y}}"):
        continue

# Add student name at the end
p_last = doc1.add_paragraph()
p_last.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
p_last.add_run("\n\n\n\n\n{{ho_ten}}").bold = True

doc1.save(os.path.join(public_dir, "1. Mau 10_KND_Ban tu kiem diem_DHKT_2025.docx"))

# ------------------------------------------------------------
# 2. Fix class dots in Biên bản họp lớp (6. Biên bản họp lớp.docx)
# ------------------------------------------------------------
print("Fixing BB Hop Lop...")
doc6 = docx.Document(os.path.join(public_dir, "6. Biên bản họp lớp.docx"))
for p in doc6.paragraphs:
    if "Lớp……………" in p.text and "đã tiến hành họp nhận xét" in p.text:
        # Re-assign paragraph to clean out dots completely
        p.text = "Lớp {{lop}} đã tiến hành họp nhận xét về những ưu điểm, khuyết điểm chính của đảng viên dự bị {{ho_ten}} trong suốt quá trình phấn đấu học tập, công tác, rèn luyện tại lớp, cụ thể như sau:"

doc6.save(os.path.join(public_dir, "6. Biên bản họp lớp.docx"))

# ------------------------------------------------------------
# 3. Fix faculty name in Biên bản Liên chi đoàn (3. Biên bản họp Liên chi Đoàn.docx)
# ------------------------------------------------------------
print("Fixing BB LCD...")
doc3 = docx.Document(os.path.join(public_dir, "3. Biên bản họp Liên chi Đoàn.docx"))
for p in doc3.paragraphs:
    if "Liên chi Đoàn khoa {{ho_ten}}" in p.text:
        p.text = p.text.replace("Liên chi Đoàn khoa {{ho_ten}}", "Liên chi Đoàn khoa {{khoa}}")

doc3.save(os.path.join(public_dir, "3. Biên bản họp Liên chi Đoàn.docx"))

print("All fixes completed successfully!")
