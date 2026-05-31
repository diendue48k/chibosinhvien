import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

templates = [
    "1. Mau 10_KND_Ban tu kiem diem_DHKT_2025.doc",
    "1. Nghị quyết LCĐ.docx",
    "1. Nghị quyết Đoàn Trường (02 bản).docx",
    "3. Biên bản họp Liên chi Đoàn.docx",
    "4. Nghị quyết Chi Đoàn.docx",
    "5. Biên bản họp Chi Đoàn.docx",
    "6. Biên bản họp lớp.docx",
    "ĐHKT_QĐ 213- 24052026.Sinh vien.docx"
]

public_dir = r"D:\CBSV\public"
scratch_dir = r"D:\CBSV\scratch"

for t in templates:
    path = os.path.join(public_dir, t)
    if not os.path.exists(path):
        continue
    
    out_path = os.path.join(scratch_dir, t.replace(".docx", "_dump.txt").replace(".doc", "_dump.txt"))
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(f"TEMPLATE DUMP FOR: {t}\n")
        f.write("=" * 80 + "\n\n")
        
        if t.endswith('.doc'):
            f.write("Binary .doc file - cannot dump via python-docx.\n")
            continue
            
        try:
            doc = docx.Document(path)
            f.write("--- PARAGRAPHS ---\n")
            for idx, p in enumerate(doc.paragraphs):
                f.write(f"P{idx}: {p.text}\n")
                
            f.write("\n--- TABLES ---\n")
            for t_idx, table in enumerate(doc.tables):
                f.write(f"\nTABLE {t_idx}:\n")
                for r_idx, row in enumerate(table.rows):
                    row_text = []
                    for c_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.replace("\n", " ").strip()
                        row_text.append(f"C{c_idx}: {cell_text}")
                    f.write(f"  ROW {r_idx}: {', '.join(row_text)}\n")
        except Exception as e:
            f.write(f"Error reading file: {e}\n")

print("Dumping completed!")
