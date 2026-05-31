import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

templates = [
    "1. Nghị quyết LCĐ.docx",
    "1. Nghị quyết Đoàn Trường (02 bản).docx",
    "3. Biên bản họp Liên chi Đoàn.docx",
    "4. Nghị quyết Chi Đoàn.docx",
    "5. Biên bản họp Chi Đoàn.docx",
    "6. Biên bản họp lớp.docx"
]

public_dir = r"D:\CBSV\public"

for t in templates:
    path = os.path.join(public_dir, t)
    if not os.path.exists(path):
        continue
    doc = docx.Document(path)
    print(f"\n==================== {t} ====================")
    for idx, p in enumerate(doc.paragraphs):
        text = p.text
        if "{{" in text:
            print(f"  P{idx}: {text}")
                
    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    text = p.text
                    if "{{" in text:
                        print(f"  Table Cell (Row {r_idx}, Col {c_idx}): {text}")
